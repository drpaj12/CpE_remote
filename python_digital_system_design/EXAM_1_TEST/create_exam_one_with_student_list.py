import sys
import os
import subprocess
import random
import operator
import csv
# DOCUMENT creater
from docx import Document
# FUNCTIONS to generate exam
from boolean_2_level_circuit_ex1_q1 import create_two_equations
from kmap_problem_create_ex1_q2 import create_kmap_question
# EMAIL STUFF
import smtplib, ssl
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders

#--------------------
def send_email(username, directory):
    # sends the email attachment
    msg = MIMEMultipart()
    sender = '' # senders email needs to be added here
    receiver = username+'@miamioh.edu' # NOTE this is a miamioh.edu address - modify to your need
    cc = "" # this can be used to document
    receivers = [cc] + [receiver]
    password = "" # password that you've setup on your email
    msg['From'] = sender
    msg['To'] = receiver
    msg['Cc'] = cc
    msg['Subject'] = "Exam 1 - ECE 287 - Attached docx - Attempt 1"
    attachment = directory+'/'+username+'.docx'
    docfile = open(attachment, "rb")
    part = MIMEBase('application', "octet-stream")
    part.set_payload(docfile.read())
    encoders.encode_base64(part)
    attachment_txt = 'attachment; filename="'+username+'.docx"'
    part.add_header('Content-Disposition', attachment_txt)
    msg.attach(part)
	
    # app password - created on gmail account
    smtpObj = smtplib.SMTP('smtp.gmail.com', 587)
    smtpObj.ehlo()
    smtpObj.starttls()
    smtpObj.login(sender, password)
	
    # actual send
    print("sender:"+sender)
    print("receivers:")
    print(*receivers)
    #print("message="+msg.as_string())
    smtpObj.sendmail( sender, receivers, msg.as_string())
	
    smtpObj.quit()

#--------------------
def create_exam_doc(directory, username, code, q1_str_o1, q1_str_o2, q2_str_equation):
     document = Document()
     
     # heading
     document.add_heading('Exam 1 - ECE 287', level=1)
     document.add_heading('User:'+username+' , code:'+str(code), level=1)
     # pargraph
     document.add_heading('Instructions', level=2)
     document.add_paragraph('- This is a take home exam.  The work must be your own.  Any common code will result in the exam being forwarded to administration for academic integrity violation.  Do not show, give, or tell anyone else how you are solving these problems.  See academic integrity in the syllabus or email me if you have any questions.')
     document.add_paragraph('-Submit a pdf or word file of your solutions.')
     document.add_paragraph('-See the rubric on canvas for grading')
     document.add_paragraph('-Please read all instructions carefully!')
     
     # page break
     document.add_page_break()
     
     # Question 1
     paragraph = document.add_paragraph('')
     run = paragraph.add_run('1) ')
     run.bold = True
     paragraph.add_run('Answer/design for these two logic equations (in Verilog):\n'+q1_str_o1+'\n'+q1_str_o2+'\n\na) Draw a schematic of this circuit\nb) What is the truth table for o1 and o2\nc) Create a complete Verilog module using a combinational always block and defining all input, output, and internal signals\nd) How many transistors in your schematic design\ne) Assuming all gates have a delay of 1ns, what is the longest path delay in your schematic (highlight this path) and report the delay')
     
     # page break
     document.add_page_break()
     
     #Question 2
     paragraph = document.add_paragraph('')
     run = paragraph.add_run('2) ')
     run.bold = True
     paragraph.add_run(' For the following logic equation :\n'+q2_str_equation+' \na) Minimize and write the boolean equation (or Verilog) using a Karnaugh Map where {a} is the top variable for each of the two k-maps and {b,c} are on the y-axis and {d,e} are on the x-axis of the maps.\nb) How many transistors in the orginal equation?\nc) How many transistors in the minimized equation from a)?')
     
     document.save(directory+'/'+username+'.docx')
#--------------------

# 0_script 1_list_of_student_usernames 2_directory_to_store_files
# SAMPLE: create_exam_one_with_student_list.py student_list.csv OUTFILES
print ("Number of arguments: %d" %  len(sys.argv))
print ("Argument List: %s" % str(sys.argv))

directory_out = sys.argv[2]
code = 0
# read through the file and run them 
with open(sys.argv[1]) as f:
    reader = csv.reader(f)
    for student_name, subject_line, body, attach in reader:
        print("Student Name: %s" % student_name)
        print("Subject Line: %s" % subject_line)
        print("Body: %s" % body)
        print("Attachment: %s" % attach)

        # question1
        o1_final_gate_out, o2_final_gate_out = create_two_equations(directory_out, student_name)

        # question2 
        q2_sum = create_kmap_question(directory_out, student_name)

        # create the exam
        create_exam_doc(sys.argv[2], student_name, code, 'o1 = '+o1_final_gate_out, 'o2 = '+o2_final_gate_out, q2_sum)

        # send the email
        send_email(student_name, sys.argv[2])

        # update the code
        code = code + 1
 
f.close()

