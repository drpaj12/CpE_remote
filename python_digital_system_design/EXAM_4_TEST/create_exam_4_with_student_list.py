import sys
import os
import subprocess
import random
import operator
import csv
from docx import Document
from selector_ex4 import *
# EMAIL STUFF
import smtplib, ssl
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders

#--------------------
def send_email(username, directory):
    # sends the email attachment
    msg = MIMEMultipart()
    sender = '' # senders email needs to be added here
    receiver = username+'@miamioh.edu' # NOTE this is a miamioh.edu address - modify to your need
    cc = "" # this can be used to document
    receivers = [cc] + [receiver]
    password = "" # password that you've setup on your email
    msg['From'] = sender
    msg['To'] = receiver
    msg['Cc'] = cc
    msg['Subject'] = "Exam 1 - ECE 287 - Attached docx - Attempt 1"
    attachment = directory+'/'+username+'.docx'
    docfile = open(attachment, "rb")
    part = MIMEBase('application', "octet-stream")
    part.set_payload(docfile.read())
    encoders.encode_base64(part)
    attachment_txt = 'attachment; filename="'+username+'.docx"'
    part.add_header('Content-Disposition', attachment_txt)
    msg.attach(part)

    # app password - created on gmail account
    smtpObj = smtplib.SMTP('smtp.gmail.com', 587)
    smtpObj.ehlo()
    smtpObj.starttls()
    smtpObj.login(sender, password)

    # actual send
    print("sender:"+sender)
    print("receivers:")
    print(*receivers)
    #print("message="+msg.as_string())
    smtpObj.sendmail( sender, receivers, msg.as_string())

    smtpObj.quit()

#--------------------
def create_exam_doc(directory, username, code, M, N):
    document = Document()
    
    # heading
    document.add_heading('Exam 4 - ECE 287', level=1)
    document.add_heading('User:'+username+' , code:'+str(code), level=1)
    # pargraph
    document.add_heading('Instructions', level=2)
    document.add_paragraph('- This is a take home exam.  The work must be your own.  Any common code will result in the exam being forwarded to administration for academic integrity violation.  Do not show, give, or tell anyone else how you are solving these problems.  See academic integrity in the syllabus or email me if you have any questions.')
    document.add_paragraph('-Submit a zip file with code as specificed below and a README.txt for additional answers and comments.')
    document.add_paragraph('-See the rubric on canvas for grading')
    document.add_paragraph('-Please read all instructions carefully!')
    
    # page break
    document.add_page_break()
    
    # Question 1
    paragraph = document.add_paragraph('')
    run = paragraph.add_run('1) ')
    run.bold = True
    paragraph.add_run('Design the following system in Verilog:\n')
    
    paragraph.add_run('You are to implement the Verilog design for your given C code (below) with a testbench.  Each exam is unique and has some of the following operations as described below (you only have a subset of these functions, but all of them are listed).\n\n')

    paragraph.add_run('You need to build your module in the provided wrapper available on Canvas with a testbench called - "tb.v".  You need to set a “done” signal when your code is complete and start on a “start signal” (see the wrapper).  Use the wrapper file as it has a test bench and reads in the inputs to start your program.  You can change the testbench “tb.v” to test your system properly (you can provide commented code for different instances and should have at least 4 test cases implemented in the testbench).  In your submission you will include the following in your zip file:\n-tb.v = your testbench with at least 4 different cases tested\n-ex4_wrapper.v = ia single file with the wrapper code and your implementation including all Verilog code (do not submit other Verilog files except the testbench)\n-memory.qip = your memory qip file named exactly like this\n-memory.v = the verilog file for your memory\n\n')

    paragraph.add_run('For your system, you will be dealing with N=' + str(N) + ' bit numbers and your memory will have an M=' + str(M) + ' number of words in it.\n')

    # page break
    document.add_page_break()
    paragraph = document.add_paragraph('')

    boldc1 = paragraph.add_run('YOUR CODE TO IMPLEMENT:\n')
    boldc1.bold = True

    paragraph.add_run('// all integers (int) are '+str(N)+' bits\n')
    paragraph.add_run('void ex4_wrapper(int in1, int in2)\n')
    paragraph.add_run('{\n')
    paragraph.add_run('\tint x, y, output1;\n')
    paragraph.add_run('\tint i, idx;\n\n')
    paragraph.add_run('\tint M = '+str(M)+';\n')
    paragraph.add_run('\tint memory[M];\n\n')
    paragraph.add_run('\tmem_init_rand(memory, M, 6);\n')
    paragraph.add_run('\tmem_display(memory, M); // output data in the memory\n\n')

    mem = random.randrange(5)
    if (mem == 0):
        paragraph.add_run('\tx = mem_min(memory, M);\n')
    elif (mem == 1):
        paragraph.add_run('\tx = mem_max(memory, M);\n')
    elif (mem == 2):
        paragraph.add_run('\tx = mem_mid(memory, M);\n')
    elif (mem == 3):
        paragraph.add_run('\tx = mem_instances_idx(memory, M, abs(in1));\n')
    else:
        paragraph.add_run('\tx = mem_instances(memory, M, abs(in1));\n')

    paragraph.add_run('\tif (x < 10)\n')
    paragraph.add_run('\t\tx = 10;\n')
    paragraph.add_run('\telse if (x >= 20)\n')
    paragraph.add_run('\t\tx=19\n\n')
    paragraph.add_run('\ty = abs(in2);\n')
    paragraph.add_run('\tif (y < 6)\n')
    paragraph.add_run('\t\ty = 6;\n')
    paragraph.add_run('\telse if (y >= 9)\n')
    paragraph.add_run('\t\ty=8\n\n')
    paragraph.add_run('\toutput1 = 0;\n')
    paragraph.add_run('\tidx = abs(rand()) % 3;\n')
    paragraph.add_run('\twhile(idx < x)\n')
    paragraph.add_run('\t{\n')
    paragraph.add_run('\t\tfor (i = '+str(random.randrange(3))+'; i < y; i++)\n')
    paragraph.add_run('\t\t{\n')

    func = random.randrange(5)
    if (func == 0):
        paragraph.add_run('\t\t\tmemory[i+idx] = pop_count(i*idx);\n')
    elif (func == 1):
        paragraph.add_run('\t\t\tmemory[i+idx] = palindrome(memory[idx]);\n')
    elif (func == 2):
        paragraph.add_run('\t\t\tmemory[i+idx] = nib_swap(memory[idx]);\n')
    elif (func == 3):
        paragraph.add_run('\t\t\tmemory[i+idx] = rot_left(memory[idx], i);\n')
    elif (func == 4):
        paragraph.add_run('\t\t\tmemory[i+idx] = rot_right(memory[idx], i);\n')
    else:
        # not used...too easy
        paragraph.add_run('\t\t\tmemory[i] = abs(memory[i*idx]);\n')

    func = random.randrange(3)
    if (func == 0):
        paragraph.add_run('\t\t\toutput1 = output1 + min(memory[i], 0-i-idx, abs(in1), in2);\n')
    elif (func == 1):
        paragraph.add_run('\t\t\toutput1 = output1 + max(memory[i], idx+i, abs(in1), in2);\n')
    else:
        paragraph.add_run('\t\t\toutput1 = output1 + add_min_max(memory[i], idx+i, abs(in1), in2);\n')

    paragraph.add_run('\t\t}\n')
    paragraph.add_run('\t\tidx += '+str(random.randrange(2)+1)+'\n')
    paragraph.add_run('\t}\n\n')
    paragraph.add_run('\tmem_display(memory, M); // output all data in the memory\n')
    paragraph.add_run('\tprintf("%d\\n", output1); // show the output \n')
    paragraph.add_run('}\n')

    boldc2 = paragraph.add_run('END YOUR CODE TO IMPLEMENT\n\n')
    boldc2.bold = True

    bold1 = paragraph.add_run('Notes:\n-Looping and conditionals must be implemented as finite state machines.\n-Do not optimize the C code.\n-Each of your functions must be implemented in a separate module with the same name.  For example, if I need to implement abs(a) then my module will be called "module abs(..."\n-You can use any basic Verilog logic or +, –, *, /, and % operaroters.  Do not use Verilog we have not learned in the class.  For example, do not use "for" or "forgen".\n- Compilation warnings, errors, and time will be looked at for grading purposes.  Avoid “inferred latch” in particular.\n- Your memory should be created using the IP for "On Chip Memory" and should be a 1-port RAM.  To allow different pieces to access that memory you will need to build a combinational multiplexer (if or case in Verilog) to access that memory.\n  \n\n')
    bold1.bold = True

    ul = paragraph.add_run('MEMORY FUNCTIONS\n')
    ul.underline = True
    paragraph.add_run('- mem_display(memory, M) = outputs the data in memory.  Implement as an FSM that iterates through the memory of size M\n- mem_init_rand(memory, M, max_val_bits) = this initializes the memory with random numbers where there are M words, of N bits, and the values will be two\'s compliment values from POSITIVE 2^(max_val_bits-1)-1 to NEGATIVE 2^(max_val_bits-1).  Use an FSM and a LFRS (included in the wrapper) to initialize the memory.  Since the max_val_bits is constant in the C code, you can solve it assuming the constant, but make sure you sign extend negative values when storing for your word size N (ie.  111111 = -1 for 6 bits and 1111 1111 1111 1111 = -1 for N=16 bits).\n- mem_max(memory, M) = this returns the value of the largest number in the memory (assuming two\'s compliment) in the memory.\n- mem_min(memory, M) = this returns the value of the smallest number in the memory.  (assuming two\'s compliment)\n-mem_mid(memory, M) = finds the arithmetic average value in the middle of the memory rounded up.\n- mem_instances_idx(memory, M, find) = this returns the index (the location in memory) of the value find in memory.  It returns "-1" if the value is not found\n- mem_instances(memory, M, find) = this returns the number of times the value "find" is located in memory.\n\n')

    ul = paragraph.add_run('MIN/MAX FUNCTIONS\n')
    ul.underline = True
    paragraph.add_run('Assume "a", "b", "c", "d" are two\'s compliment numbers that are N bits\n- min(a, b, c, d) = assuming the 4 inputs are two\'s compliment, this function returns the smallest value of the four values.\n- max(a, b, c, d) = assuming the 4 inputs are two\'s compliment, this function returns the largest value of the four values.\n- add_min_max(a, b, c, d) = assuming the 4 inputs are two\'s compliment, this function returns the sum of the largest value plus the smallest.\n\n')

    ul = paragraph.add_run('DATA FUNCTIONS\n')
    ul.underline = True
    paragraph.add_run('Assume "a" is a two\'s compliment numbers that are N bits and "y" is a positive value that is big enough to work in your system\n- random() - returns a random integer of N bits.\n- pop_count(a) - returns how many "1s" are in the binary value a.\n- palindrome(a) - returns the palindrome (invert the values around a mirror in the center).  For example, palindrome(16’b1000100100000010) returns 16’b0100000010010001.\n- nib_swap(a) = swaps the least significant nibble (4-bits) with the most significant nibble. For example, nib_swap(16’b0000111100111100) returns 16’b1100111100110000.\n- abs(a) = absolute value takes the absolute value of a number.  For example, abs(-3) returns 3.\nrot_left(a,y) – rotates x by a bits in the left direction.  Rotate left means that the most significant bit gets put into the least significant spot and all other bits shift 1 to the left.  For example, rot_left(10001000,1) for an 8 bit number returns 00010001.\nrot_right(a,y) – rotates x by a bits in the right direction.  Rotate right means that the least significant bit gets put into the most significant spot and all other bits shift 1 to the right per rotation.  For example, rot_right(10001001,1) for an 8 bit number returns 11000100.\n\n')
     
    bl = paragraph.add_run('TIPs\n')
    bl.bold = True
    paragraph.add_run('- Each of the modules that has an internal FSM needs a start and done signal to communicate back with your high-level FSM.\n- Start small and test frequently.  This is called a spiral design approach, and this will work well in this case\n')
 
    # page break
    document.add_page_break()
    
    # Question 2
    paragraph = document.add_paragraph('')
    run = paragraph.add_run('2) ')
    run.bold = True
    paragraph.add_run('For the above design, report the FPGA area results (LE, mem bits, DSP, etc.) and for one of your test cases, how long does it take to execute on a DE2-115 with 50MHz clock?\n')

    document.save(directory+'/'+username+'.docx')
#--------------------

# 0_script 1_list_of_student_usernames 2_directory_to_store_files
# SAMPLE: create_exam_4_with_student_list.py student_list.csv OUTFILES
print ("Number of arguments: %d" %  len(sys.argv))
print ("Argument List: %s" % str(sys.argv))

directory_out = sys.argv[2]
code = 0
# read through the file and run them 
with open(sys.argv[1]) as f:
    reader = csv.reader(f)
    for student_name, subject_line, body, attach in reader:
        print("Student Name: %s" % student_name)
        print("Subject Line: %s" % subject_line)
        print("Body: %s" % body)
        print("Attachment: %s" % attach)

        # get parameters
        word_size,number_of_words = pick_mem_params()
        print (str(word_size)+" "+str(number_of_words))

        # create the exam
        create_exam_doc(sys.argv[2], student_name, code, number_of_words, word_size)

        # send the email
        send_email(student_name, sys.argv[2])

        # update the code
        code = code + 1
 
f.close()

