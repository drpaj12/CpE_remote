import sys
import os
import subprocess
import random
import operator
import csv
from docx import Document
from equation_for_4_inputs_ex2_q1 import create_equation
# EMAIL STUFF
import smtplib, ssl
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders

#--------------------
def send_email(username, directory):
    # sends the email attachment
    msg = MIMEMultipart()
    sender = '' # senders email needs to be added here
    receiver = username+'@miamioh.edu' # NOTE this is a miamioh.edu address - modify to your need
    cc = "" # this can be used to document
    receivers = [cc] + [receiver]
    password = "" # password that you've setup on your email
    msg['From'] = sender
    msg['To'] = receiver
    msg['Cc'] = cc
    msg['Subject'] = "Exam 1 - ECE 287 - Attached docx - Attempt 1"
    attachment = directory+'/'+username+'.docx'
    docfile = open(attachment, "rb")
    part = MIMEBase('application', "octet-stream")
    part.set_payload(docfile.read())
    encoders.encode_base64(part)
    attachment_txt = 'attachment; filename="'+username+'.docx"'
    part.add_header('Content-Disposition', attachment_txt)
    msg.attach(part)

    # app password - created on gmail account
    smtpObj = smtplib.SMTP('smtp.gmail.com', 587)
    smtpObj.ehlo()
    smtpObj.starttls()
    smtpObj.login(sender, password)

    # actual send
    print("sender:"+sender)
    print("receivers:")
    print(*receivers)
    #print("message="+msg.as_string())
    smtpObj.sendmail( sender, receivers, msg.as_string())

    smtpObj.quit()

#--------------------
def create_exam_doc(directory, username, code, q1_str_eqn, q1_str_values):
     document = Document()
     
     # heading
     document.add_heading('Exam 2 - ECE 287', level=1)
     document.add_heading('User:'+username+' , code:'+str(code), level=1)
     # pargraph
     document.add_heading('Instructions', level=2)
     document.add_paragraph('- This is a take home exam.  The work must be your own.  Any common code will result in the exam being forwarded to administration for academic integrity violation.  Do not show, give, or tell anyone else how you are solving these problems.  See academic integrity in the syllabus or email me if you have any questions.')
     document.add_paragraph('-Submit a pdf or word file of your solutions.')
     document.add_paragraph('-See the rubric on canvas for grading')
     document.add_paragraph('-Please read all instructions carefully!')
     
     # page break
     document.add_page_break()
     
     # Question 1
     paragraph = document.add_paragraph('')
     run = paragraph.add_run('1) ')
     run.bold = True
     paragraph.add_run('For the following equation: out = '+q1_str_eqn+'\na) Create a schematic design.  Assume that all input signals {A, B, C, D} are 4 bit signals and are two\'s complement numbers.\nIn the schematic you are allowed to use black boxes for Full Adders (FA) and Half Adders (HA).  You can also use 2:1 mux, 4:1 mux, etc.  Also, keep the schematic simple by using black boxes and muxes and not optimized circuitry.\nb) How many transistors in your schematic (use a table to show your calculations clearly)?\nc) Given the following decimal values for the inputs ('+q1_str_values+'), show the result of the equation (in binary and decimal) and all the intermediate steps in binary.\n')
     
     # page break
     document.add_page_break()
     
     #Question 2
     paragraph = document.add_paragraph('')
     run = paragraph.add_run('2) ')
     run.bold = True
     paragraph.add_run('a) For the equation from question 1 implement a Verilog design using the the same input signals (4bits, Two\'s Compliment).  You can use the +, -, <, >, >=, <= operators in your design and you should implement the design in an always block.\nNOTE: I would make sure this compiles in Quartus and works correctly.\nb) Draw the schematic of your Verilog design and you are allowed to use black boxes for the comparison circuit (assume it has the same delay and transistor count as question 1), Full Adders (FA), and Half Adders (HA). You can also use 2:1 mux, 4:1 mux, etc.  Do not optimize the circuitry and keep it simple.\nNOTE: you must build the schematic with low-level components and you can not use Quartus RTL netlist viewer for this.  You can define a clear box, show the clear box design, and then use the clear box in your larger design.')
     
     document.save(directory+'/'+username+'.docx')
#--------------------

# 0_script 1_list_of_student_usernames 2_directory_to_store_files
# SAMPLE: create_exam_2_with_student_list.py student_list.csv OUTFILES
print ("Number of arguments: %d" %  len(sys.argv))
print ("Argument List: %s" % str(sys.argv))

directory_out = sys.argv[2]
code = 0
# read through the file and run them 
with open(sys.argv[1]) as f:
    reader = csv.reader(f)
    for student_name, subject_line, body, attach in reader:
        print("Student Name: %s" % student_name)
        print("Subject Line: %s" % subject_line)
        print("Body: %s" % body)
        print("Attachment: %s" % attach)

        # question1
        eq1,val1 = create_equation()
        print (eq1)

        # create the exam
        create_exam_doc(sys.argv[2], student_name, code, eq1, val1)

        # send the email
        send_email(student_name, sys.argv[2])

        # update the code
        code = code + 1
 
f.close()

