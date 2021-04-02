import sys
import os
import subprocess
import random
import operator
import csv
from docx import Document
from equation_for_2_sequence_ex3_q1 import create_sequence
# EMAIL STUFF
import smtplib, ssl
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders

#--------------------
def send_email(username, directory):
    # sends the email attachment
    msg = MIMEMultipart()
    sender = '' # senders email needs to be added here
    receiver = username+'@miamioh.edu' # NOTE this is a miamioh.edu address - modify to your need
    cc = "" # this can be used to document
    receivers = [cc] + [receiver]
    password = "" # password that you've setup on your email
    msg['From'] = sender
    msg['To'] = receiver
    msg['Cc'] = cc
    msg['Subject'] = "Exam 1 - ECE 287 - Attached docx - Attempt 1"
    attachment = directory+'/'+username+'.docx'
    docfile = open(attachment, "rb")
    part = MIMEBase('application', "octet-stream")
    part.set_payload(docfile.read())
    encoders.encode_base64(part)
    attachment_txt = 'attachment; filename="'+username+'.docx"'
    part.add_header('Content-Disposition', attachment_txt)
    msg.attach(part)

    # app password - created on gmail account
    smtpObj = smtplib.SMTP('smtp.gmail.com', 587)
    smtpObj.ehlo()
    smtpObj.starttls()
    smtpObj.login(sender, password)

    # actual send
    print("sender:"+sender)
    print("receivers:")
    print(*receivers)
    #print("message="+msg.as_string())
    smtpObj.sendmail( sender, receivers, msg.as_string())

    smtpObj.quit()

#--------------------
def create_exam_doc(directory, username, code, q1_seq1, q1_seq2):
     document = Document()
     
     # heading
     document.add_heading('Exam 3 - ECE 287', level=1)
     document.add_heading('User:'+username+' , code:'+str(code), level=1)
     # pargraph
     document.add_heading('Instructions', level=2)
     document.add_paragraph('- This is a take home exam.  The work must be your own.  Any common code will result in the exam being forwarded to administration for academic integrity violation.  Do not show, give, or tell anyone else how you are solving these problems.  See academic integrity in the syllabus or email me if you have any questions.')
     document.add_paragraph('-Submit a pdf or word file of your solutions.')
     document.add_paragraph('-See the rubric on canvas for grading')
     document.add_paragraph('-Please read all instructions carefully!')
     
     # page break
     document.add_page_break()
     
     # Question 1
     paragraph = document.add_paragraph('')
     run = paragraph.add_run('1) ')
     run.bold = True
     paragraph.add_run('Design the following FSM (noting you can not use shift registers to solve the problem) as both schematic and Verilog:\n- The input is a one-bit input (labeled in).\n- The output turns on (out = “1”) when the sequences “'+q1_seq1+'” or “'+q1_seq2+'” are detected. (Assume that a new bit arrives every clock cycle in a sequence on the "in" port, which is a 1-bit input)\na) Build a state diagram, corresponding state table, and schematic circuit of the above.\nb) Build a Verilog implementation of the state diagram from (a), with a rst signal, in the form:\n\nmodule fsm_question_1(rst, clk, in, out);\n input rst, clk; // normal signals for sequential control\n input in;\n output out;\n ...\n endmodule\n')
     
     # page break
     document.add_page_break()
     
     #Question 2
     paragraph = document.add_paragraph('')
     run = paragraph.add_run('2) ')
     run.bold = True
     paragraph.add_run(' In this question you have the finite state machine you designed in Verilog from the previous question.\n Also, you have a Linear Feedback Shift Register (LFSR) that generates a pseudorandom number on each clock cycle.  Note that the srand input needs to be set to some constant of 16bits:\n \n /* LFSR */\n module lfsr(clk, rst, srand, init_srand, rand);\n input clk, rst;\n input [15:0] srand;\n input init_srand;\n output reg [15:0] rand;\n always @(posedge clk or negedge rst)\n begin\n \tif (rst == 1\'b0)\n \t\trand <= 16\'d0a;\n\telse\n\t\tif(init_srand == 1\'b1)\n\t\t\trand <= srand; //This is a constant to initialize it\n\t\telse\n\t\t\t{rand[14:0],rand[12]^rand[15]^rand[3]^rand[6]};\n\nend\n \n endmodule\n \n You need to build a circuit that has X outputs (each output is a bus-signal of 64bits labelled outseqbit0, outseqbit1, ...), that tells you how many times the sequences you identify from question 1 (with your fsm_question_1 design) have occurred in a defined time period for the X least significant bits of a random number sequence generated by an LFSR.  X=3 if the first letter of your first name is A through M and X=4 if your first the first letter of your first name is N through Z.  For example, Peter is "P" which means X=4 for me.\n \n The time period of the clock is equal to Y seconds where Y is equal to the number of letters in your first and last name.  For example, “Peter Jamieson” is 13 letters, and so my time interval is 13seconds (yours will be different depending on your name length).  We will assume the clock speed of the system is 2KHz.\n \n')
     paragraph.add_run('The circuit starts when a “start” signal is switched to “1” (note we have normal asynchronous reset procedure).  After Y seconds the X outputs will appear with the respective counts.  The circuit will then wait (and maintain the calculated results on the four outputs) for the “start” signal to go low (“0”) and can repeat a new count when “start” goes high again.  This control needs to be implemented as a Finite State Machine in Verilog.  Also, I have provided the start of this Verilog module.  You must add and fill in the details to implement the above.\n \n\n module rand_num_bit_counter(\n input clk,\n input rst,\n input start,\n // Define the outputs as needed where there are X of these \n output [63:0] outseqbit0;\n output [63:0] outseqbit1;\n ...\n);\n \n ...\n \n /* other definitions */\n \n reg [    :    ] random_num;\n lfsr my_lfsr(        ,        , 16’hA1F2,      , random_num);\n \n reg [     :      ] outs_from_q1_fsms;\n \n /* instantiate fsm_question_1(s) as needed - X of them */\n fsm_question_1 bit0(               ,                   ,                   , outs_from_q1_fsm[0]);\n ...\n \n /* other reg and wires you need to do things */\n \n /* define the size of the State and Next State signals\n reg [    :    ] S; // State\n reg [    :    ] NS; // Next State\n \n /* define the state parameters for states as needed */\n parameter\n \n ...\n ')
     
     document.save(directory+'/'+username+'.docx')
#--------------------

# 0_script 1_list_of_student_usernames 2_directory_to_store_files
# SAMPLE: create_exam_2_with_student_list.py student_list.csv OUTFILES
print ("Number of arguments: %d" %  len(sys.argv))
print ("Argument List: %s" % str(sys.argv))

directory_out = sys.argv[2]
code = 0
# read through the file and run them 
with open(sys.argv[1]) as f:
    reader = csv.reader(f)
    for student_name, subject_line, body, attach in reader:
        print("Student Name: %s" % student_name)
        print("Subject Line: %s" % subject_line)
        print("Body: %s" % body)
        print("Attachment: %s" % attach)

        # question1
        seq1,seq2 = create_sequence()
        print (seq1+" "+seq2)

        # create the exam
        create_exam_doc(sys.argv[2], student_name, code, seq1, seq2)

        # send the email
        send_email(student_name, sys.argv[2])

        # update the code
        code = code + 1
 
f.close()

